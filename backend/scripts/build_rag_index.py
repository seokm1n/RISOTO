"""sources/의 원문을 청크로 잘라 대응 원칙 보충용 벡터 색인을 만든다.

    cd backend
    python -m scripts.build_rag_index                 # 전체 재색인
    python -m scripts.build_rag_index --dry-run       # 임베딩 없이 청크만 확인

색인 전용 스크립트라 런타임 의존성이 아니다. 실행에만 다음이 필요하다.

    pip install pypdf python-docx

**문서를 유형에 붙이는 방식**: principles_data.json의 sources에 어느 문서가 어느 세부
유형의 근거인지 이미 적혀 있다. 그 매핑을 그대로 재사용하므로, 자료가 추가되면
principles_data.json의 sources에 등록하는 것만으로 색인 대상이 된다.

**청크에 사용 제약을 심는 이유**: 리콜 자료는 미국 CPSC 절차를, 사이버 자료는 NIST
프레임워크 용어를 본문에 담고 있다. 청크만 떼어 프롬프트에 넣으면 국내 기업 보고서에
미국 절차가 실린다. principles_data.json의 caution을 청크마다 붙여 그 사고를 막는다.

**원본은 저장소에 없다**: sources/*는 .gitignore로 제외돼 있다(용량 158MB+). 재색인이
필요하면 sources/README.md의 출처 목록을 보고 원문을 확보해야 한다.
"""
from __future__ import annotations

import argparse
import json
import re
import sys
from pathlib import Path

ENGINE = Path("app/services/response_engine")
RAG_DIR = ENGINE / "sources"
DATA = ENGINE / "principles_data.json"


def _load_env() -> None:
    """저장소 루트의 .env를 읽어 온다.

    컨테이너 안에서는 compose가 환경변수를 넣어 주지만, 이 스크립트는 backend/에서 직접
    실행하는 용도라 그 경로를 타지 않는다. .env가 한 단계 위에 있어서 그냥 두면
    Settings가 database_url 없다고 죽는다 - 임베딩 호출 직전에야 터져서 원인을 찾기 어렵다.
    """
    env = Path("../.env")
    if not env.exists():
        return
    try:
        from dotenv import load_dotenv
    except ImportError:
        return
    load_dotenv(env)

# 한 청크의 목표 길이(자). 너무 짧으면 맥락이 끊기고, 너무 길면 프롬프트를 잡아먹는다.
TARGET = 900
OVERLAP = 150
MIN_CHUNK = 200

SUFFIXES = (".pdf", ".docx")


def _load_pypdf():
    try:
        from pypdf import PdfReader  # noqa

        return PdfReader
    except ImportError:
        print("pypdf가 필요합니다:  pip install pypdf", file=sys.stderr)
        raise


def read_pdf(path: Path) -> list[str]:
    """페이지별 텍스트. 스캔 PDF는 텍스트 레이어가 없어 빈 문자열만 나온다."""
    PdfReader = _load_pypdf()
    return [p.extract_text() or "" for p in PdfReader(str(path)).pages]


def read_docx(path: Path) -> list[str]:
    """docx는 페이지 개념이 없어 본문 전체를 한 덩어리로 돌려준다.

    표를 빼면 안 된다. 이 형식의 자료는 대응 전략 비교표·사례 비교표처럼 가장 밀도 높은
    내용이 표에 들어 있어서, 문단만 읽으면 정작 쓸모 있는 부분이 통째로 빠진다.
    표는 파이프로 구분한 행 단위 텍스트로 펴서 본문 뒤에 붙인다.
    """
    try:
        import docx
    except ImportError:
        print("python-docx가 필요합니다:  pip install python-docx", file=sys.stderr)
        raise

    d = docx.Document(str(path))
    parts = [p.text.strip() for p in d.paragraphs if p.text.strip()]
    for table in d.tables:
        rows = []
        for row in table.rows:
            cells = [c.text.strip().replace("\n", " ") for c in row.cells]
            if any(cells):
                rows.append(" | ".join(cells))
        if rows:
            parts.append("\n".join(rows))
    return ["\n\n".join(parts)]


def read_pages(path: Path) -> list[str]:
    return read_docx(path) if path.suffix.lower() == ".docx" else read_pdf(path)


def doc_type_map() -> tuple[dict[str, list[str]], dict[str, str], dict[str, str], dict[str, str]]:
    """principles_data.json에서 (문서명 -> 유형들), caution, 검증표시, 고정 파일명."""
    d = json.loads(DATA.read_text(encoding="utf-8"))
    types: dict[str, list[str]] = {}
    cautions: dict[str, str] = {}
    verifications: dict[str, str] = {}
    pinned: dict[str, str] = {}

    def note(src: dict) -> None:
        if src.get("file"):
            pinned[src["doc"]] = src["file"]
        if src.get("verification"):
            verifications[src["doc"]] = src["verification"]

    for code, entry in d["types"].items():
        for src in entry.get("sources", []):
            doc = src["doc"]
            types.setdefault(doc, [])
            if code not in types[doc]:
                types[doc].append(code)
            if entry.get("caution"):
                cautions[doc] = entry["caution"]
            note(src)
    # 공통 베이스 자료는 모든 유형에 붙인다.
    all_codes = list(d["types"])
    for src in d.get("common_base", {}).get("sources", []):
        types[src["doc"]] = all_codes
        note(src)
    return types, cautions, verifications, pinned


def _norm(s: str) -> str:
    return re.sub(r"[^\w가-힣]", "", s).lower()


def match_score(doc_name: str, path: Path) -> int:
    """문서명과 파일명의 최장 공통 부분문자열 길이. 유형 접두사를 뗀 제목을 우선 본다."""
    title = re.split(r"[：:]", doc_name)[-1]
    target = _norm(path.stem)
    best = 0
    for key in (_norm(title), _norm(doc_name)):
        if not key:
            continue
        for size in range(min(len(key), len(target)), 4, -1):
            if any(key[i : i + size] in target for i in range(len(key) - size + 1)):
                best = max(best, size)
                break
    return best


def assign_files(
    doc_names: list[str],
    files: list[Path],
    pinned: dict[str, str] | None = None,
    min_score: int = 6,
) -> dict[str, Path]:
    """문서명 -> 파일을 1:1로 배정한다. `file`이 지정된 출처가 항상 우선한다.

    **이름 유사도만으로는 위험하다**: 실제로 "State of Crisis Communication (Coombs, 2014)"이
    제목에 같은 어구가 든 CERC_Crisis_Communication_Plans.pdf를 가져가는 바람에, CERC 본문이
    Coombs 논문 이름으로 164청크 색인되고 진짜 논문(CoombsFinalWES.pdf)은 통째로 빠졌다.
    인용 출처가 뒤바뀌므로 보고서 신뢰도에 직접 영향을 준다. 그래서 principles_data.json의
    각 source에 `file`을 적어 고정하고, 유사도 매칭은 미지정 항목의 보조 수단으로만 둔다.

    나머지는 모든 조합의 점수를 구한 뒤 높은 것부터 확정하고, 한 번 쓰인 파일과 문서명은
    후보에서 뺀다. 법령 원문처럼 애초에 파일이 없는 출처는 배정되지 않고 남는다 - 정상이다.
    """
    by_name = {p.name: p for p in files}
    assigned: dict[str, Path] = {}
    used: set[Path] = set()
    for doc, fname in (pinned or {}).items():
        path = by_name.get(fname)
        if path is None:
            print(f"경고: '{doc}'에 지정된 파일 {fname}을 찾을 수 없습니다.", file=sys.stderr)
            continue
        assigned[doc] = path
        used.add(path)

    pairs = sorted(
        ((match_score(d, p), d, p) for d in doc_names for p in files),
        key=lambda t: -t[0],
    )
    for score, doc, path in pairs:
        if score < min_score or doc in assigned or path in used:
            continue
        assigned[doc] = path
        used.add(path)
    return assigned


def chunk_pages(pages: list[str]) -> list[tuple[int, str]]:
    """페이지를 이어 붙이되 TARGET 길이로 자르고 OVERLAP만큼 겹친다.

    페이지 경계로만 자르면 한 지침이 두 페이지에 걸쳐 끊기고, 무시하고 자르면 어느
    페이지에서 왔는지 잃는다. 이어 붙이면서 시작 페이지를 기록하는 방식으로 절충한다.
    """
    out: list[tuple[int, str]] = []
    buf, start_page = "", 1
    for i, text in enumerate(pages, 1):
        text = re.sub(r"\n{3,}", "\n\n", (text or "").strip())
        if not text:
            continue
        if not buf:
            start_page = i
        buf += ("\n" if buf else "") + text
        while len(buf) >= TARGET:
            out.append((start_page, buf[:TARGET]))
            buf = buf[TARGET - OVERLAP :]
            start_page = i
    if len(buf) >= MIN_CHUNK:
        out.append((start_page, buf))
    return out


def build(dry_run: bool = False) -> None:
    _load_env()
    if not DATA.exists():
        print(f"{DATA}를 찾을 수 없습니다. backend/ 에서 실행하세요.", file=sys.stderr)
        raise SystemExit(1)

    types, cautions, verifications, pinned = doc_type_map()
    files = sorted(p for p in RAG_DIR.glob("*") if p.suffix.lower() in SUFFIXES)
    if not files:
        print(f"{RAG_DIR}에 원문이 없습니다. sources/README.md의 출처 목록을 참고하세요.")
        return

    from app.services.response_engine.rag.store import Chunk, VectorStore

    assigned = assign_files(list(types), files, pinned)
    chunks: list[Chunk] = []
    matched, unmatched, empty = [], [], []
    for doc_name, codes in types.items():
        path = assigned.get(doc_name)
        if path is None:
            unmatched.append(doc_name)
            continue
        pieces = chunk_pages(read_pages(path))
        if not pieces:
            # 스캔 PDF는 텍스트 레이어가 없어 여기로 떨어진다. 조용히 넘기면
            # 색인에서 빠진 사실을 모르게 되므로 따로 보고한다.
            empty.append((doc_name, path.name))
            continue
        matched.append((doc_name, path.name, codes))
        for j, (page, text) in enumerate(pieces):
            chunks.append(Chunk(
                chunk_id=f"{path.stem[:24]}#{j:03d}",
                text=text,
                doc=doc_name,
                page=page,
                risk_types=list(codes),
                caution=cautions.get(doc_name),
                verification=verifications.get(doc_name),
            ))

    print(f"매칭된 문서 {len(matched)}개 / 청크 {len(chunks)}개")
    for doc, fname, codes in matched:
        n = sum(1 for c in chunks if c.doc == doc)
        print(f"  {n:4d}청크  {','.join(codes):<28} {fname[:44]}")
    if empty:
        print(f"\n텍스트를 못 뽑은 파일 {len(empty)}개 (스캔본 여부 확인 필요):")
        for doc, fname in empty:
            print(f"  - {fname[:60]}")
    if unmatched:
        print(f"\n파일을 못 찾은 출처 {len(unmatched)}개 (principles_data.json의 doc 표기 확인 필요):")
        for u in unmatched:
            print("  -", u[:70])

    if dry_run:
        print("\n--dry-run: 임베딩 없이 종료")
        return
    if not chunks:
        print("색인할 청크가 없습니다.")
        return

    from app.services.response_engine.rag.embed import MODEL, embed

    print(f"\n임베딩 중... ({MODEL}, {len(chunks)}청크)")
    vectors, usage = embed([c.text for c in chunks])
    store = VectorStore(chunks, vectors, MODEL)
    store.save()
    from app.services.response_engine.rag.store import DEFAULT_DIR

    print(f"저장 -> {DEFAULT_DIR}")
    print(f"토큰 {usage['tokens']:,} / 호출 {usage['calls']}회")


def parse_args():
    p = argparse.ArgumentParser(description="대응 원칙 보충용 RAG 색인 구축")
    p.add_argument("--dry-run", action="store_true", help="임베딩 없이 청크 구성만 확인")
    return p.parse_args()


if __name__ == "__main__":
    build(parse_args().dry_run)
